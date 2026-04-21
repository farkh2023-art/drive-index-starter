import pytest
from pathlib import Path
from docx import Document
from PIL import Image


@pytest.fixture
def sample_docx(tmp_path):
    path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Ceci est un document test.")
    doc.add_paragraph("Deuxieme paragraphe.")
    doc.save(path)
    return path


@pytest.fixture
def sample_image(tmp_path):
    path = tmp_path / "test.png"
    img = Image.new("RGB", (800, 600), color=(200, 100, 50))
    img.save(path)
    return path


@pytest.fixture
def sample_txt(tmp_path):
    path = tmp_path / "test.txt"
    path.write_text("Contenu texte de test.\nDeuxieme ligne.", encoding="utf-8")
    return path


@pytest.fixture
def minimal_template(tmp_path):
    path = tmp_path / "index_template.html"
    path.write_text(
        "<html><script>\n"
        "// Cette partie sera remplie automatiquement par le script Python\n"
        "    // ... (ne rien changer ici)\n"
        "</script></html>",
        encoding="utf-8",
    )
    return path


@pytest.fixture
def sample_dir(tmp_path, sample_docx, sample_image, sample_txt):
    """Directory with mixed file types for integration tests."""
    subdir = tmp_path / "sous_dossier"
    subdir.mkdir()
    (subdir / "note.txt").write_text("Note dans sous-dossier.", encoding="utf-8")
    return tmp_path
